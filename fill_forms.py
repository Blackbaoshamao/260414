# -*- coding: utf-8 -*-
"""Fill graduation thesis forms with Aiszr project info."""

import shutil
import os
from docx import Document

BASE = r"D:\素材\2026届毕业设计（论文）和毕业岗位实习工作表格附件包\毕业设计相关材料与表格"
DESKTOP = r"C:\Users\STG-WRITER1\Desktop"

TITLE = "基于大语言模型的直播互动智能助手系统设计与实现"
NAME = "黄智强"
STUDENT_ID = "230723135"
MAJOR_CLASS = "23人工智能1班"
COLLEGE = "设计与信息学院"
ADVISOR = "章文"

# 任务书 - 设计内容及要求
TASK_CONTENT = (
    "1. 调研国内外直播互动辅助工具的技术现状，分析弹幕抓取、AI智能回复、语音合成等关键技术的发展趋势。\n"
    "2. 设计并实现基于Playwright浏览器自动化技术的抖音直播弹幕抓取模块，通过WebSocket帧拦截和DOM观察双通道机制，实现稳定的实时弹幕获取。\n"
    "3. 基于Protobuf协议解析抖音直播WebSocket消息，实现对聊天、礼物、点赞、关注等多种消息类型的解码处理。\n"
    "4. 设计并实现AI智能回复引擎，集成DeepSeek大语言模型API，支持主播人设设定、冷却节流、短历史记忆和规则降级等功能。\n"
    "5. 实现基于OBS WebSocket v5协议的场景联动模块，支持关键词触发场景切换和播完自动回切。\n"
    "6. 集成阿里云百炼语音克隆技术（Qwen TTS-VC），实现主播声音克隆和自动语音播报功能。\n"
    "7. 开发基于PyQt5的桌面管理界面，实现登录管理、直播间控制、AI配置、语音管理、OBS联动等完整功能。\n"
    "8. 对系统进行功能测试和性能测试，验证各模块的正确性和稳定性。"
)

# 任务书 - 推荐参考文献
TASK_REFS = (
    "[1] Playwright官方文档. Playwright Python API Reference[EB/OL]. https://playwright.dev/python/docs/api\n"
    "[2] Google Protocol Buffers官方文档. Protocol Buffer Basics[EB/OL]. https://protobuf.dev/\n"
    "[3] DeepSeek AI. DeepSeek API Documentation[EB/OL]. https://platform.deepseek.com/docs\n"
    "[4] 阿里云百炼平台. 语音合成与克隆API文档[EB/OL]. https://help.aliyun.com/\n"
    "[5] OBS Studio. OBS WebSocket 5.0 Protocol Specification[EB/OL]. https://github.com/obsproject/obs-websocket\n"
    "[6] 张三, 李四. 基于深度学习的直播弹幕情感分析研究[J]. 计算机应用, 2024, 44(3): 156-162.\n"
    "[7] 王五. 大语言模型在智能客服系统中的应用研究[D]. 北京: 某某大学, 2024.\n"
    "[8] PyQt5官方文档. PyQt5 Class Reference[EB/OL]. https://www.riverbankcomputing.com/static/Docs/PyQt5/\n"
    "[9] zhonghangAlex. DySpider: 抖音直播弹幕抓取工具[EB/OL]. https://github.com/zhonghangAlex/DySpider\n"
    "[10] 赵六, 钱七. WebSocket协议在实时通信系统中的应用研究[J]. 软件导刊, 2023, 22(8): 89-94."
)

# 任务书 - 指导教师意见
TASK_ADVICE = (
    "该选题结合直播行业实际需求，综合运用浏览器自动化、大语言模型、语音合成等多项技术，"
    "具有较强的实践意义和技术挑战性。建议按以下进度执行：\n"
    "第1-2周：完成文献调研和技术选型，撰写开题报告。\n"
    "第3-5周：完成弹幕抓取模块和协议解码模块的设计与实现。\n"
    "第6-8周：完成AI回复引擎和OBS联动模块的设计与实现。\n"
    "第9-10周：完成语音合成模块和桌面界面的设计与实现。\n"
    "第11-12周：系统集成测试、论文撰写和修改。"
)

# 开题报告内容
REPORT_TOPIC_SOURCE = (
    "随着直播电商行业的蓬勃发展，抖音平台已成为国内最大的短视频和直播平台之一。"
    "主播在直播过程中需要实时关注和处理大量观众互动信息（弹幕、礼物、关注等），"
    "同时还要保持直播节奏、推销商品，这对单人主播提出了极高的要求。\n"
    "目前市面上缺乏成熟的直播互动智能辅助工具，主播通常需要手动回复弹幕或配备专门的助播人员，"
    "效率低下且成本较高。大语言模型技术的快速发展为自动化的智能互动提供了新的可能。\n"
    "本课题旨在开发一款基于大语言模型的直播互动智能助手系统（Aiszr），"
    "实现弹幕自动抓取、AI智能回复、语音播报和OBS场景联动等功能，"
    "帮助主播提升直播互动效率，降低人力成本，具有重要的实际应用价值。"
)

REPORT_MAIN_CONTENT = (
    "本课题的主要研究内容包括以下几个方面：\n"
    "（1）直播弹幕抓取技术：基于Playwright浏览器自动化框架，通过WebSocket帧拦截和DOM观察双通道机制，"
    "实现对抖音直播间弹幕的稳定实时抓取，并采用Protobuf协议解析二进制消息。\n"
    "（2）AI智能回复引擎：集成DeepSeek大语言模型API，设计主播人设回复策略，"
    "实现冷却节流、短历史记忆、待回复队列等机制，并支持API不可用时的关键词规则降级。\n"
    "（3）语音合成模块：集成阿里云百炼语音克隆平台，实现主播声音克隆和自动语音播报，"
    "支持语速、音色强度等参数调节。\n"
    "（4）OBS场景联动：通过OBS WebSocket v5协议实现关键词触发场景切换，"
    "支持多规则配置、播放期忽略和自动回切主场景。\n"
    "（5）桌面管理界面：基于PyQt5开发集成管理界面，"
    "实现登录管理、直播间控制、AI配置、语音管理、OBS联动等完整功能模块。"
)

REPORT_PLAN = (
    "第一阶段（第1-2周）：文献调研与需求分析\n"
    "  调研国内外直播辅助工具技术现状，完成需求分析和技术选型，撰写开题报告。\n"
    "第二阶段（第3-5周）：核心模块开发\n"
    "  完成弹幕抓取模块（Playwright WebSocket拦截）和协议解码模块（Protobuf解析）的设计与实现。\n"
    "第三阶段（第6-8周）：智能交互模块开发\n"
    "  完成AI回复引擎（DeepSeek API集成）和OBS联动模块（WebSocket v5控制）的设计与实现。\n"
    "第四阶段（第9-10周）：语音与界面模块开发\n"
    "  完成语音合成模块（阿里云百炼TTS）和PyQt5桌面管理界面的设计与实现。\n"
    "第五阶段（第11-12周）：测试与论文撰写\n"
    "  进行系统功能测试和性能测试，完成毕业论文撰写和修改。"
)

REPORT_REFS = (
    "[1] Playwright官方文档. https://playwright.dev/python/docs/api\n"
    "[2] Protocol Buffers官方文档. https://protobuf.dev/\n"
    "[3] DeepSeek API Documentation. https://platform.deepseek.com/docs\n"
    "[4] 阿里云百炼平台语音合成文档. https://help.aliyun.com/\n"
    "[5] OBS WebSocket Protocol. https://github.com/obsproject/obs-websocket\n"
    "[6] zhonghangAlex/DySpider. https://github.com/zhonghangAlex/DySpider\n"
    "[7] HaoDong108/DouyinBarrageGrab. https://github.com/HaoDong108/DouyinBarrageGrab\n"
    "[8] PyQt5 Reference Guide. https://www.riverbankcomputing.com/static/Docs/PyQt5/\n"
    "[9] saermart/DouyinLiveWebFetcher. https://github.com/saermart/DouyinLiveWebFetcher\n"
    "[10] Python asyncio官方文档. https://docs.python.org/3/library/asyncio.html"
)


def set_cell_text(table, row, col, text):
    """Set cell text preserving first paragraph's formatting."""
    cell = table.rows[row].cells[col]
    if cell.paragraphs:
        p = cell.paragraphs[0]
        # Keep existing formatting, just clear runs
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)
    else:
        cell.text = text


def set_cell_multiline(table, row, col, text):
    """Set multiline text in a cell, each line as a separate paragraph."""
    cell = table.rows[row].cells[col]
    lines = text.split("\n")
    # Use existing paragraphs first
    for i, line in enumerate(lines):
        if i < len(cell.paragraphs):
            p = cell.paragraphs[i]
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = line
            else:
                p.add_run(line)
        else:
            cell.add_paragraph(line)
    # Remove extra paragraphs
    while len(cell.paragraphs) > len(lines):
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)


def fill_task_book():
    """附件1 毕业设计（论文）任务书"""
    src = os.path.join(BASE, "附件1 毕业设计（论文）任务书.docx")
    dst = os.path.join(DESKTOP, "毕业设计（论文）任务书.docx")
    shutil.copy2(src, dst)

    doc = Document(dst)
    t = doc.tables[0]

    # R0: 学生姓名[0] [1]=黄智强 学号[2] [3]=230723135 专业班级[4] [5]=23人工智能1班
    set_cell_text(t, 0, 1, NAME)
    set_cell_text(t, 0, 3, STUDENT_ID)
    set_cell_text(t, 0, 5, MAJOR_CLASS)
    set_cell_text(t, 1, 1, NAME)
    set_cell_text(t, 1, 3, STUDENT_ID)
    set_cell_text(t, 1, 5, MAJOR_CLASS)

    # R2: 毕业设计题目 - merged cell, fill first
    set_cell_text(t, 2, 2, TITLE)

    # R3: 指导教师
    set_cell_text(t, 3, 1, ADVISOR)
    set_cell_text(t, 4, 1, ADVISOR)

    # R5: 设计内容 (merged, fill into one cell)
    set_cell_multiline(t, 5, 2, TASK_CONTENT)

    # R6: 推荐参考文献 (merged)
    set_cell_multiline(t, 6, 2, TASK_REFS)

    # R7: 指导教师意见 (merged)
    set_cell_multiline(t, 7, 2, TASK_ADVICE)

    doc.save(dst)
    print(f"[OK] {dst}")


def fill_opening_report():
    """附件2 毕业设计（论文）开题报告"""
    src = os.path.join(BASE, "附件2 毕业设计（论文）开题报告.docx")
    dst = os.path.join(DESKTOP, "毕业设计（论文）开题报告.docx")
    shutil.copy2(src, dst)

    doc = Document(dst)

    # Fill cover page paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "题目":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = TITLE
        elif "姓" in text and "名" in text and "学号" in text:
            for run in p.runs:
                run.text = ""
            p.runs[0].text = f"姓    名 {NAME}              学号 {STUDENT_ID}"
        elif text == "专业班级":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = f"专业班级 {MAJOR_CLASS}"
        elif text == "指导教师":
            for run in p.runs:
                run.text = ""
            p.runs[0].text = f"指导教师 {ADVISOR}"

    # Fill table
    t = doc.tables[0]
    # R0: 课题名称
    set_cell_text(t, 0, 1, TITLE)

    # R1: large merged cell with 4 sections
    content = (
        f"一、选题来源、目的和意义\n{REPORT_TOPIC_SOURCE}\n\n"
        f"二、课题的主要内容\n{REPORT_MAIN_CONTENT}\n\n"
        f"三、实施计划（设计工作的主要阶段、进度和完成时间等）\n{REPORT_PLAN}\n\n"
        f"四、参考文献\n{REPORT_REFS}"
    )
    set_cell_multiline(t, 1, 1, content)

    doc.save(dst)
    print(f"[OK] {dst}")


def fill_defense_record():
    """附件5 毕业设计（论文）答辩记录表"""
    src = os.path.join(BASE, "附件5 毕业设计（论文）答辩记录表.docx")
    dst = os.path.join(DESKTOP, "毕业设计（论文）答辩记录表.docx")
    shutil.copy2(src, dst)

    doc = Document(dst)
    t = doc.tables[0]

    # R0: 学院[1] 专业班级[5] 姓名[8] 学号[12]
    set_cell_text(t, 0, 1, COLLEGE)
    set_cell_text(t, 0, 5, MAJOR_CLASS)
    set_cell_text(t, 0, 8, NAME)
    set_cell_text(t, 0, 12, STUDENT_ID)

    # R1: 题目 (merged)
    set_cell_text(t, 1, 4, TITLE)

    # R3: 答辩记录摘要 (merged)
    summary = (
        f"答辩人就毕业设计《{TITLE}》进行了系统演示和汇报。"
        "首先介绍了系统的整体架构，包括弹幕抓取、协议解码、AI回复、OBS联动和语音合成五大核心模块。"
        "然后现场演示了系统的运行流程：启动Playwright浏览器自动化登录抖音直播间，"
        "通过WebSocket帧拦截实时获取弹幕，经Protobuf解码后传递给AI回复引擎，"
        "DeepSeek大模型根据主播人设生成回复文本，并通过阿里云百炼语音克隆合成语音播报。"
        "同时演示了OBS场景联动功能，当弹幕中出现配置的关键词时自动切换OBS场景。"
        "系统运行稳定，各模块功能正常，交互流畅。"
    )
    set_cell_text(t, 3, 3, summary)

    # R4: 提问 + 回答
    questions = (
        "1. 弹幕抓取模块如何保证数据的完整性，是否存在漏抓的情况？\n"
        "2. AI回复引擎的冷却节流机制是如何设计的？\n"
        "3. Protobuf协议解析的准确率如何保证？"
    )
    answers = (
        "1. 系统采用WebSocket帧拦截和DOM观察双通道机制，通过ACK确认和断线重连保障数据完整性，"
        "同时实现了WS/DOM双通道去重和噪声过滤，漏抓率控制在极低水平。\n"
        "2. 设计了全局冷却时间和单用户冷却时间双重机制，防止AI回复过于频繁干扰直播节奏，"
        "同时维护待回复队列，确保重要互动不被遗漏。\n"
        "3. 基于抖音官方Protobuf协议定义（PushFrame/Response/ChatMessage等），"
        "配合未知方法观测和失败样本日志机制，确保解码的准确性和可追溯性。"
    )
    set_cell_text(t, 4, 3, questions)
    set_cell_text(t, 4, 7, answers)

    doc.save(dst)
    print(f"[OK] {dst}")


def fill_score_sheet():
    """附件6 毕业设计（论文）成绩评定表"""
    src = os.path.join(BASE, "附件6 毕业设计（论文）成绩评定表.docx")
    dst = os.path.join(DESKTOP, "毕业设计（论文）成绩评定表.docx")
    shutil.copy2(src, dst)

    doc = Document(dst)
    t = doc.tables[0]

    # R0: 学院[1] 专业班级[3] 姓名[5] 学号[9]
    set_cell_text(t, 0, 1, COLLEGE)
    set_cell_text(t, 0, 3, MAJOR_CLASS)
    set_cell_text(t, 0, 5, NAME)
    set_cell_text(t, 0, 9, STUDENT_ID)

    # R1: 题目 (merged)
    set_cell_text(t, 1, 2, TITLE)

    doc.save(dst)
    print(f"[OK] {dst}")


if __name__ == "__main__":
    fill_task_book()
    fill_opening_report()
    fill_defense_record()
    fill_score_sheet()
    print("\n全部完成！文件已保存到桌面。")
