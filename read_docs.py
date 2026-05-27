"""Extract text + structural info from the three docx files."""
import os
from docx import Document
from docx.shared import Pt

files = {
    "template": r"D:\论文\附件4 毕业设计（论文）参考样本.docx",
    "spec": r"D:\论文\附件3 毕业设计（论文）撰写规范及格式.docx",
    "thesis": r"D:\论文\毕业材料\毕业材料\毕业论文_基于大语言模型的直播互动智能助手系统设计与实现_附件4模板版.docx",
}

def summarize(path, out_path, max_paragraphs=None):
    doc = Document(path)
    lines = []
    lines.append(f"=== FILE: {path}")
    lines.append(f"=== Section count: {len(doc.sections)}")
    for i, sec in enumerate(doc.sections):
        lines.append(f"[Section {i}] page H={sec.page_height} W={sec.page_width} "
                     f"margins T={sec.top_margin} B={sec.bottom_margin} L={sec.left_margin} R={sec.right_margin}")
    lines.append(f"=== Paragraph count: {len(doc.paragraphs)}")
    lines.append("")

    for idx, p in enumerate(doc.paragraphs):
        if max_paragraphs and idx >= max_paragraphs:
            lines.append(f"... ({len(doc.paragraphs) - idx} more paragraphs truncated)")
            break
        text = p.text
        style = p.style.name if p.style else "?"
        align = str(p.alignment) if p.alignment is not None else "-"
        pf = p.paragraph_format
        line_sp = pf.line_spacing
        # first run formatting
        first_run = None
        for r in p.runs:
            if r.text.strip():
                first_run = r
                break
        if first_run is None and p.runs:
            first_run = p.runs[0]
        if first_run is not None:
            font = first_run.font
            size = font.size.pt if font.size else "-"
            bold = font.bold
            name = font.name
        else:
            size = bold = name = "-"
        lines.append(f"[{idx:03d}] style={style!r} align={align} ls={line_sp} font={name} size={size} bold={bold}")
        if text.strip():
            lines.append(f"     TEXT: {text[:200]}")
        # tables not in paragraphs; show separately
        lines.append("")

    # tables
    lines.append(f"=== Table count: {len(doc.tables)}")
    for ti, t in enumerate(doc.tables):
        lines.append(f"[Table {ti}] rows={len(t.rows)} cols={len(t.columns)}")
        for ri, row in enumerate(t.rows[:5]):
            cells = [c.text.replace('\n', ' | ')[:60] for c in row.cells]
            lines.append(f"   row{ri}: {cells}")

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"wrote {out_path}")

os.makedirs(r"d:\Pjt\260414\thesis_extracted", exist_ok=True)
summarize(files["template"], r"d:\Pjt\260414\thesis_extracted\template.txt")
summarize(files["spec"], r"d:\Pjt\260414\thesis_extracted\spec.txt")
summarize(files["thesis"], r"d:\Pjt\260414\thesis_extracted\thesis.txt")
