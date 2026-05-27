"""Build thesis by editing 附件4 in place: keep cover & TOC styles untouched,
swap in the real content.

Strategy
--------
1. Open 附件4 directly.
2. Keep cover paragraphs (index 0..22) and the single section-break paragraph
   right after them (index 23). Cover is left exactly as-is.
3. Delete every paragraph / table after that, up to but NOT including the
   final <w:sectPr> at the bottom of the body.
4. Append: abstract (template "Plain Text" title + "Normal"), TOC (template
   "toc 1" / "toc 2"), then body (template "Heading 1" / "Heading 2" / "Normal"),
   page breaks between chapters via pageBreakBefore.
5. Save as 毕业论文_..._最终格式版.docx.

Why this preserves formatting
-----------------------------
We apply *named styles* defined inside 附件4 itself. python-docx writes
<w:pStyle w:val="..."/>, so the run-level fonts / sizes / line-spacing that
the template encodes in styles.xml carry over automatically.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FIG_DIR = r"d:\Pjt\260414\thesis_figs"
TEMPLATE = r"D:\论文\附件4 毕业设计（论文）参考样本.docx"
OUT_DIR = r"D:\论文\毕业材料\毕业材料"
OUT_NAME = "毕业论文_基于大语言模型的直播互动智能助手系统设计与实现_最终格式版.docx"
OUT_PATH = os.path.join(OUT_DIR, OUT_NAME)

# Styles we rely on from 附件4 (论文模板 group)
STYLE_NORMAL    = "Normal"
STYLE_H1        = "Heading 1"
STYLE_H2        = "Heading 2"
STYLE_PLAINTEXT = "Plain Text"
STYLE_TOC1      = "toc 1"
STYLE_TOC2      = "toc 2"
STYLE_FIGCAP    = "Normal (Web)"   # template uses this style for figure captions

FONT_SONG = "宋体"
FONT_HEI  = "黑体"

# ------------------------------------------------------------------ helpers

def set_run_cn_font(run, font_name, size_pt=None, bold=None):
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

def add_p(doc, text="", style=None, align=None, *, font_name=None, size_pt=None,
          bold=None, page_break_before=False, line_spacing=None,
          space_before=None, space_after=None, first_line_indent_chars=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if page_break_before:
        p.paragraph_format.page_break_before = True
    if first_line_indent_chars is not None:
        pPr = p._element.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind"); pPr.append(ind)
        ind.set(qn("w:firstLineChars"), str(first_line_indent_chars * 100))
        ind.set(qn("w:firstLine"), str(int(first_line_indent_chars * 240)))
    if text:
        run = p.add_run(text)
        if font_name is None:
            font_name = FONT_SONG
        set_run_cn_font(run, font_name, size_pt=size_pt, bold=bold)
    return p

def add_body(doc, text):
    return add_p(doc, text, style=STYLE_NORMAL, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 font_name=FONT_SONG, size_pt=12, line_spacing=1.5,
                 first_line_indent_chars=2)

def add_h1(doc, text, page_break_before=True):
    return add_p(doc, text, style=STYLE_H1, align=WD_ALIGN_PARAGRAPH.CENTER,
                 font_name=FONT_SONG, size_pt=15, bold=True,
                 line_spacing=1.5, space_before=24, space_after=12,
                 page_break_before=page_break_before)

def add_h2(doc, text):
    return add_p(doc, text, style=STYLE_H2, align=WD_ALIGN_PARAGRAPH.LEFT,
                 font_name=FONT_SONG, size_pt=12, bold=True,
                 line_spacing=1.5, space_before=12, space_after=6)

def add_centered(doc, text, font_name, size_pt, bold=False, style=None,
                 space_before=0, space_after=0):
    return add_p(doc, text, style=style, align=WD_ALIGN_PARAGRAPH.CENTER,
                 font_name=font_name, size_pt=size_pt, bold=bold,
                 line_spacing=1.5, space_before=space_before, space_after=space_after)

def add_figure(doc, image_path, caption, width_inch=5.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(image_path, width=Inches(width_inch))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.5
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(caption)
    set_run_cn_font(run, FONT_SONG, size_pt=10.5)

def add_table_caption(doc, caption):
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.5
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.keep_with_next = True
    run = cap.add_run(caption)
    set_run_cn_font(run, FONT_SONG, size_pt=10.5)

def _set_table_borders(table):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tblPr.append(borders)

def add_table(doc, headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Template may not define "Table Grid"; apply borders by XML directly.
    try:
        t.style = doc.styles["Table Grid"]
    except KeyError:
        _set_table_borders(t)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths_cm:
        for ci, cw in enumerate(col_widths_cm):
            for ri in range(1 + len(rows)):
                t.cell(ri, ci).width = Cm(cw)

    def fill(cell, text, bold=False):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.25
        run = para.add_run(text)
        set_run_cn_font(run, FONT_SONG, size_pt=10.5, bold=bold)

    for ci, h in enumerate(headers):
        fill(t.rows[0].cells[ci], h, bold=True)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            fill(t.rows[ri].cells[ci], val)

# ------------------------------------------------------------------ strip placeholder

def strip_after_cover(doc, keep_through_p_index: int):
    """Delete every paragraph / table after the first `keep_through_p_index + 1`
    paragraphs, up to but not including the final <w:sectPr> of body.
    """
    body = doc.element.body
    children = list(body)

    # Walk children, count paragraphs, find the cutoff position
    p_idx = -1
    cut = None
    for i, c in enumerate(children):
        tag = c.tag.split("}")[-1]
        if tag == "p":
            p_idx += 1
            if p_idx == keep_through_p_index:
                cut = i + 1   # delete starting from i+1
                break

    if cut is None:
        raise RuntimeError("could not locate cover boundary paragraph")

    # Delete from `cut` upwards but skip the final top-level sectPr
    final_sectPr = None
    if children[-1].tag.endswith("}sectPr"):
        final_sectPr = children[-1]

    for c in children[cut:]:
        if c is final_sectPr:
            continue
        body.remove(c)

# ------------------------------------------------------------------ build

def build():
    doc = Document(TEMPLATE)

    # Cover (paragraphs 0..22) + the section-break paragraph at 23 stay intact.
    # Paragraph 23's <w:sectPr> ends the cover section; we keep that boundary
    # so the cover page geometry remains untouched.
    strip_after_cover(doc, keep_through_p_index=23)

    # ---------------- ABSTRACT ----------------
    # Title line (same style the template uses for the abstract heading)
    add_centered(doc, "基于大语言模型的直播互动智能助手系统设计与实现",
                 FONT_HEI, 18, bold=True, style=STYLE_PLAINTEXT,
                 space_before=12, space_after=18)

    # 摘要
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run("【摘要】"); set_run_cn_font(r1, FONT_HEI, 12, bold=True)
    r2 = p.add_run(
        "随着直播电商行业的蓬勃发展，直播间已成为商品销售和品牌推广的重要渠道。"
        "然而，主播在直播过程中面临着弹幕信息过载、互动回复不及时、多平台工具切换繁琐等问题。"
        "本文设计并实现了一个基于大语言模型的直播互动智能助手系统（Aiszr），"
        "集成了弹幕实时抓取、AI智能回复、OBS场景联动和AI语音播报等核心功能。"
        "系统采用Python技术栈，利用Playwright浏览器自动化技术实现抖音直播间弹幕的"
        "WebSocket帧拦截与Protobuf协议解码，并通过DOM观察通道提供补偿，"
        "实现了高完整率的弹幕采集；接入DeepSeek大语言模型并通过人设系统、回复节流、"
        "用户冷却与规则降级等多层策略保证AI回复质量与节奏；"
        "基于OBS WebSocket v5实现关键词触发的场景联动；"
        "基于阿里云百炼实现语音克隆与TTS合成，并通过ffmpeg + HLS将绿幕视频与TTS音频合成后"
        "推流到OBS，形成完整的数字人直播辅助闭环。"
        "实测表明，系统在4小时连续直播场景下运行稳定，"
        "弹幕端到端延迟低于1秒，AI回复端到端延迟控制在5至8秒，"
        "能够显著降低主播的操作负担、提升直播间互动效率。")
    set_run_cn_font(r2, FONT_SONG, 12)

    # 关键词
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run("【关键词】"); set_run_cn_font(r1, FONT_HEI, 12, bold=True)
    r2 = p.add_run("大语言模型；直播互动；弹幕抓取；智能回复；语音克隆；OBS联动")
    set_run_cn_font(r2, FONT_SONG, 12)

    # ---------------- TOC ----------------
    # Page break to start TOC on a new page
    pb = doc.add_paragraph()
    pb.add_run().add_break(WD_BREAK.PAGE)

    # TOC title — same look as template ("目  录（论文模板）") but we drop the
    # "（论文模板）" suffix.
    add_centered(doc, "目　　录", FONT_HEI, 16, bold=True,
                 space_before=12, space_after=12)

    TOC = [
        (1, "1 绪论", 1),
        (2, "1.1 课题的研究背景", 1),
        (2, "1.2 国内外研究现状", 1),
        (2, "1.3 研究目标与内容", 2),
        (1, "2 相关技术概述", 3),
        (2, "2.1 Playwright 浏览器自动化技术", 3),
        (2, "2.2 Protobuf 协议与 WebSocket 通信", 3),
        (2, "2.3 大语言模型与 DeepSeek API", 4),
        (2, "2.4 阿里云百炼语音克隆技术", 4),
        (1, "3 系统需求分析", 5),
        (2, "3.1 功能需求分析", 5),
        (2, "3.2 非功能需求分析", 5),
        (1, "4 系统设计", 6),
        (2, "4.1 系统总体架构", 6),
        (2, "4.2 弹幕抓取模块设计", 8),
        (2, "4.3 AI 智能回复引擎设计", 10),
        (2, "4.4 语音合成模块设计", 12),
        (2, "4.5 模块依赖与关键参数", 13),
        (1, "5 系统实现", 15),
        (2, "5.1 登录会话管理与弹幕抓取", 15),
        (2, "5.2 协议解码与消息处理", 17),
        (2, "5.3 AI 智能回复引擎实现", 19),
        (2, "5.4 OBS 联动与数字人推流", 21),
        (2, "5.5 开发问题与解决方案", 24),
        (1, "6 系统测试", 28),
        (2, "6.1 功能测试", 28),
        (2, "6.2 性能测试", 29),
        (1, "结语", 31),
        (1, "参考文献", 32),
        (1, "致谢", 33),
    ]
    for level, name, page in TOC:
        style = STYLE_TOC1 if level == 1 else STYLE_TOC2
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        # Right tab stop with dot leader for page number alignment
        p.paragraph_format.tab_stops.add_tab_stop(
            Mm(155), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        run = p.add_run(name + "\t" + str(page))
        set_run_cn_font(run, FONT_SONG, 12)

    # ---------------- BODY ----------------
    # CHAPTER 1
    add_h1(doc, "1 绪论")
    add_h2(doc, "1.1 课题的研究背景")
    add_body(doc,
        "近年来，随着移动互联网技术的快速发展和5G网络的普及，直播电商行业迎来了"
        "爆发式增长。据商务部数据显示，2024年上半年全国直播电商销售额达到1.27万亿元，"
        "同比增长超过30%。抖音、快手、淘宝直播等平台已经成为商品销售的重要渠道，"
        "越来越多的商家和品牌选择通过直播方式进行产品推广和销售。")
    add_body(doc,
        "在直播场景中，主播需要同时处理多个任务：展示商品、回答观众提问、管理直播间氛围、"
        "控制OBS场景切换等。这种多任务并行的工作模式对主播的注意力和精力提出了很高的要求。"
        "特别是在人气较高的直播间，弹幕消息大量涌入，主播往往无法及时回复所有观众的提问，"
        "导致互动体验下降，进而影响销售转化率。")
    add_body(doc,
        "为了解决上述问题，业界开始探索利用人工智能技术辅助直播运营。"
        "基于大语言模型（LLM）的智能回复系统能够自动理解观众提问并生成高质量回复，"
        "语音合成技术可以将文字回复转化为语音播报，"
        "OBS联动技术可以实现关键词触发的场景自动切换。"
        "这些技术的综合应用有望显著降低主播的工作负担，提升直播间的互动效率。")
    add_body(doc,
        "因此，本课题旨在设计并实现一个基于大语言模型的直播互动智能助手系统，"
        "通过技术手段解决直播场景中的弹幕信息过载、互动回复不及时、工具切换繁琐等痛点，"
        "为直播电商行业提供一套完整的智能化解决方案。")

    add_h2(doc, "1.2 国内外研究现状")
    add_body(doc,
        "在弹幕抓取方面，目前市面上的主要方案包括直接WebSocket连接和浏览器自动化两种方式。"
        "直接连接方案（如DouyinLiveWebFetcher、DySpider等项目）通过逆向分析抖音的"
        "WebSocket协议实现直连，但需要持续维护签名算法，且协议变更频繁导致连接不稳定。"
        "浏览器自动化方案（基于Selenium或Playwright）则利用浏览器自身的WebSocket连接"
        "进行被动监听，无需逆向签名算法，连接稳定性更高。本系统采用Playwright浏览器自动化"
        "方案，并通过DOM观察作为补偿通道，有效提升了弹幕抓取的完整性。")
    add_body(doc,
        "在AI智能回复方面，随着ChatGPT、DeepSeek等大语言模型的出现和普及，"
        "基于LLM的对话系统在客服、教育、娱乐等领域得到了广泛应用。"
        "在直播场景中，已有部分产品尝试接入大语言模型实现自动回复，"
        "但普遍存在回复内容与主播人设不符、回复频率过高导致刷屏、缺乏上下文理解等问题。"
        "本系统通过设计多层回复策略（人设系统、节流控制、用户冷却、规则降级）"
        "有效解决了这些问题。")
    add_body(doc,
        "在语音合成方面，近年来语音克隆技术取得了显著进展。阿里云百炼平台提供了"
        "基于Qwen模型的语音克隆服务，只需15秒的语音样本即可实现高质量的语音复刻，"
        "为直播场景的语音播报提供了技术基础。")
    add_body(doc,
        "在OBS联动方面，OBS Studio提供了WebSocket v5接口，允许外部程序通过"
        "JSON-RPC协议控制场景切换、媒体播放等操作，为本系统的场景联动功能"
        "提供了技术支撑。")

    add_h2(doc, "1.3 研究目标与内容")
    add_body(doc,
        "本课题的研究目标是设计并实现一个功能完整、运行稳定的直播互动智能助手系统，"
        "主要研究内容包括：")
    add_body(doc,
        "（1）研究基于Playwright的直播间弹幕被动式抓取技术，"
        "实现WebSocket帧拦截、Protobuf协议解码和DOM观察补偿的双通道弹幕采集方案。")
    add_body(doc,
        "（2）研究基于大语言模型的直播间智能回复策略，"
        "设计人设系统、回复节流、用户冷却和规则降级等多层控制机制，"
        "实现高质量的自动互动回复。")
    add_body(doc,
        "（3）研究OBS WebSocket协议的场景联动技术，"
        "实现基于关键词匹配和命中计数的自动场景切换方案。")
    add_body(doc,
        "（4）研究基于阿里云百炼的语音克隆与合成技术，"
        "实现AI回复的语音播报和数字人推流功能。")

    # CHAPTER 2
    add_h1(doc, "2 相关技术概述")
    add_h2(doc, "2.1 Playwright 浏览器自动化技术")
    add_body(doc,
        "Playwright是微软开发的浏览器自动化框架，支持Chromium、Firefox和WebKit"
        "三大浏览器引擎。与传统的Selenium相比，Playwright具有原生WebSocket监听能力，"
        "可以通过page.on(\"websocket\")事件捕获浏览器中所有WebSocket连接的数据帧，"
        "无需额外的网络代理或中间人攻击。")
    add_body(doc,
        "本系统利用Playwright的persistent context（持久化浏览器上下文）功能，"
        "将浏览器状态（包括Cookie、localStorage等）持久化到本地磁盘，"
        "实现一次扫码登录、多次复用的效果。同时，Playwright的自动等待机制和事件驱动架构"
        "使得弹幕抓取过程无需轮询，降低了CPU占用和响应延迟。")
    add_body(doc,
        "在反检测方面，Playwright启动时添加了"
        "\"--disable-blink-features=AutomationControlled\"参数，"
        "移除了浏览器的自动化标识特征，降低了被抖音平台检测为自动化工具的风险。")

    add_h2(doc, "2.2 Protobuf 协议与 WebSocket 通信")
    add_body(doc,
        "抖音直播间的弹幕数据通过WebSocket协议传输，消息体采用Protobuf（Protocol Buffers）"
        "序列化格式。Protobuf是Google开发的高效二进制序列化协议，"
        "具有体积小、解析快、跨语言兼容等优点。")
    add_body(doc,
        "抖音的弹幕传输采用三层协议结构：外层为PushFrame信封，包含序列号、日志ID等元信息；"
        "中间层为gzip压缩的负载数据；内层为Response消息，"
        "包含一个或多个具体的弹幕消息（如ChatMessage、GiftMessage、LikeMessage等）。")
    add_body(doc,
        "本系统通过手工定义Protobuf消息结构（PushFrame、Response、ChatMessage等），"
        "利用google-protobuf库的动态消息工厂（MessageFactory）"
        "实现了无需protoc编译的协议解码方案。"
        "这种方式既保持了协议解析的正确性，又避免了编译环境的依赖问题。")

    add_h2(doc, "2.3 大语言模型与 DeepSeek API")
    add_body(doc,
        "大语言模型（Large Language Model，LLM）是基于Transformer架构的深度学习模型，"
        "通过在海量文本数据上进行预训练，获得了强大的自然语言理解和生成能力。"
        "DeepSeek是由深度求索公司开发的开源大语言模型，"
        "提供了与OpenAI API兼容的接口，支持对话补全、函数调用等功能。")
    add_body(doc,
        "本系统通过HTTP请求调用DeepSeek的Chat Completions接口，"
        "将观众弹幕作为用户消息、主播人设作为系统提示词发送给模型，"
        "获取AI生成的回复文本。接口调用采用异步方式（httpx.AsyncClient），"
        "避免阻塞主线程。同时，系统设置了请求超时机制（6至12秒自适应），"
        "防止单次回复耗时过长影响整体互动节奏。")
    add_body(doc,
        "在人设系统设计方面，本系统定义了角色名称、角色定位、语气风格、回复策略、"
        "场景话术、回复限制和禁忌事项七个维度的人设配置，"
        "通过组合这些维度生成完整的系统提示词（System Prompt），"
        "确保AI回复的内容和风格与主播的人设保持一致。")

    add_h2(doc, "2.4 阿里云百炼语音克隆技术")
    add_body(doc,
        "阿里云百炼（DashScope）平台提供了基于Qwen模型的语音克隆和语音合成服务。"
        "语音克隆（Voice Cloning）技术允许用户上传15秒的语音样本，"
        "系统据此生成一个个性化的语音模型，后续的文字转语音（TTS）操作"
        "即可使用该模型生成具有用户音色特征的语音。")
    add_body(doc,
        "本系统使用阿里云百炼的Qwen Voice Enrollment模型完成声音注册，"
        "使用Qwen3-TTS-VC模型进行语音合成。合成过程支持语速和音量参数调节，"
        "合成结果为WAV格式的音频文件。系统还实现了合成缓存机制，"
        "基于文本内容、模型、语速等参数的哈希值作为缓存键，"
        "避免相同内容重复调用API，节省成本和响应时间。")
    add_body(doc,
        "在本地播放方面，系统针对Windows平台使用winsound进行异步WAV播放，"
        "针对其他平台使用pygame.mixer作为播放后端，保证了跨平台的兼容性。")

    # CHAPTER 3
    add_h1(doc, "3 系统需求分析")
    add_h2(doc, "3.1 功能需求分析")
    add_body(doc, "通过对直播场景的深入调研和分析，本系统确定了以下功能需求：")
    add_body(doc,
        "（1）登录会话管理：支持Playwright扫码登录，浏览器状态持久化，"
        "登录有效性检测和过期自动重登。")
    add_body(doc,
        "（2）弹幕实时抓取：通过WebSocket拦截和DOM观察的双通道方案，"
        "实时抓取直播间的弹幕消息（chat）、礼物通知（gift）、点赞（like）、"
        "关注（follow）、进入直播间（enter）和在线统计（stats）等消息类型。")
    add_body(doc,
        "（3）AI智能回复：接入DeepSeek大语言模型，根据弹幕内容自动生成"
        "符合主播人设的回复文本，支持回复节流（最小回复间隔）、"
        "用户冷却（同一用户冷却期）、待回复队列和规则降级"
        "（LLM不可用时回退到规则引擎）。")
    add_body(doc,
        "（4）OBS场景联动：通过OBS WebSocket v5协议，实现关键词匹配触发的"
        "场景自动切换，支持命中计数窗口、最小命中次数、规则冷却和全局冷却等控制参数，"
        "播完后自动切回主场景。")
    add_body(doc,
        "（5）AI语音播报：接入阿里云百炼语音克隆服务，"
        "支持主播和助播两个角色的声音管理，实现AI回复文本到语音的自动播报。")
    add_body(doc,
        "（6）数字人推流：集成ffmpeg实现绿幕数字人视频与TTS音频的合成，"
        "通过HLS协议推流到OBS进行播放。")

    add_h2(doc, "3.2 非功能需求分析")
    add_body(doc,
        "（1）稳定性：系统应能在长时间直播（4小时以上）中稳定运行，"
        "不出现内存泄漏、线程死锁或连接异常中断等问题。"
        "弹幕抓取的漏抓率应控制在5%以内。")
    add_body(doc,
        "（2）实时性：弹幕从产生到UI展示的端到端延迟不超过1秒，"
        "AI回复从弹幕接收到回复生成的延迟不超过5秒。")
    add_body(doc,
        "（3）易用性：系统采用PyQt5桌面应用形式，提供图形化的配置界面，"
        "用户无需编写代码即可完成所有功能的配置和使用。")
    add_body(doc,
        "（4）安全性：所有API密钥和凭据信息本地存储，不通过网络传输；"
        "API密钥输入框支持密文显示切换。")

    # CHAPTER 4
    add_h1(doc, "4 系统设计")
    add_h2(doc, "4.1 系统总体架构")
    add_body(doc,
        "本系统采用分层模块化架构设计，整体结构分为四个层次："
        "用户界面层、业务逻辑层、数据采集层和外部服务层。"
        "图4-1展示了系统的总体架构。")
    add_body(doc,
        "用户界面层基于PyQt5和PyQt-SiliconUI组件库构建，"
        "包含首页（系统状态总览）、直播页（弹幕展示与AI回复）、AI配置页、"
        "语音设置页、OBS联动页、数字人推流页和通用设置页共七个功能页面。")
    add_body(doc,
        "业务逻辑层由CaptureWorker（捕获工作器）作为核心调度中心，"
        "在独立的QThread中运行asyncio事件循环，"
        "协调弹幕处理、AI回复、OBS联动、TTS语音和数字人推流等异步任务。"
        "CaptureWorker通过PyQt信号（Signal）与UI层通信，"
        "通过asyncio.run_coroutine_threadsafe桥接Qt主线程和asyncio工作线程。")
    add_body(doc,
        "数据采集层由RoomCapture（房间捕获器）和DanmakuDecoder（弹幕解码器）组成，"
        "负责通过Playwright拦截直播间的WebSocket帧并解码为结构化消息。")
    add_body(doc,
        "外部服务层包括DeepSeek API（大语言模型）、阿里云百炼API（语音服务）、"
        "OBS WebSocket（场景控制）和ffmpeg（视频合成），"
        "通过异步HTTP或WebSocket协议与这些服务进行交互。")
    add_figure(doc, os.path.join(FIG_DIR, "fig4_1_architecture.png"),
               "图4-1 系统总体架构图")

    add_table_caption(doc, "表4-1 核心模块职责表")
    add_table(doc,
        headers=["模块名称", "主要职责", "实现说明"],
        rows=[
            ["AiszrApp", "桌面应用入口与页面路由",
             "创建首页、直播页、AI 配置页、语音页、OBS 页、数字人页和设置页。"],
            ["CaptureWorker", "异步调度中心",
             "在 QThread 中运行 asyncio 事件循环，协调弹幕、AI、语音和 OBS 任务。"],
            ["RoomCapture", "直播间采集",
             "通过 Playwright 监听 WebSocket 帧，并用 DOM 观察作为补偿通道。"],
            ["DanmakuDecoder", "协议解码",
             "解析 PushFrame、gzip 负载和 Response 消息，输出扁平化弹幕事件。"],
            ["AIReplyEngine", "AI 回复决策与生成",
             "执行人设、节流、冷却、降级与待回复队列，调用 DeepSeek 接口。"],
            ["ObsActionCtrl", "OBS 场景联动",
             "通过 OBS WebSocket v5 实现关键词命中触发与冷却。"],
            ["TTSWorker / VoiceManager", "语音合成与播放",
             "异步队列调度 TTS 合成、本地 winsound/pygame 播放、合成缓存。"],
            ["DigitalHumanPipeline", "数字人推流流水线",
             "ffmpeg 合成 HLS 切片，启动本地 HTTP 服务器并配置 OBS 媒体源。"],
        ],
        col_widths_cm=[3.2, 4.5, 7.8])

    add_h2(doc, "4.2 弹幕抓取模块设计")
    add_body(doc,
        "弹幕抓取模块采用WebSocket拦截和DOM观察的双通道方案，"
        "确保弹幕采集的完整性。整体流程如图4-2所示。")
    add_body(doc,
        "WebSocket通道：通过Playwright的page.on(\"websocket\")事件监听"
        "浏览器中建立的WebSocket连接，筛选包含\"douyin.com\"的URL，"
        "通过ws.on(\"framereceived\")事件捕获每帧数据。"
        "接收到的二进制帧经过三层解码（PushFrame解析、gzip解压、Response解析），"
        "提取出结构化的弹幕消息。系统还会发送ACK心跳包保持连接活跃。")
    add_body(doc,
        "DOM观察通道：通过注入JavaScript的MutationObserver"
        "监听直播间聊天区域的DOM变化，"
        "当检测到新的聊天消息DOM节点时，提取昵称和内容文本，"
        "通过Playwright的expose_binding回调传递给Python端处理。"
        "DOM观察作为WebSocket通道的补充，"
        "特别在WebSocket帧解析失败时提供兜底保障。")
    add_body(doc,
        "双通道融合：系统对两条通道的消息进行去重处理。"
        "每条chat消息通过（昵称, 内容）的标准化签名进行标识，"
        "在2秒的时间窗口内，同一签名只保留最先到达的那条。"
        "WebSocket通道的消息优先级高于DOM通道，当WS消息完整时直接转发，"
        "DOM消息仅在WS消息缺失昵称时作为补充。")
    add_figure(doc, os.path.join(FIG_DIR, "fig4_2_capture_flow.png"),
               "图4-2 弹幕采集与双通道融合流程图")

    add_h2(doc, "4.3 AI 智能回复引擎设计")
    add_body(doc,
        "AI智能回复引擎是本系统的核心创新点，"
        "采用了多层控制策略确保回复质量和互动节奏。"
        "回复决策的整体流程如图4-3所示。")
    add_body(doc,
        "人设系统：系统定义了角色名称、角色定位、语气风格、回复策略、场景话术、"
        "回复限制和禁忌事项七个维度的人设配置。"
        "这些配置被组合成完整的System Prompt，"
        "指导大语言模型生成符合主播人设的回复内容。"
        "例如，系统可以配置\"慵懒松弛\"的语气风格、40至80字的字数限制，"
        "以及禁止鸡汤文学、客服腔调等禁忌事项。")
    add_body(doc,
        "回复节流：系统设置了最小回复间隔（默认30秒），"
        "两次自动回复之间至少间隔设定时间，避免AI回复过于频繁导致刷屏。"
        "同时设置了每用户冷却（默认60秒），"
        "同一用户在冷却期内不会触发新的AI回复，"
        "确保回复机会均匀分配给不同观众。")
    add_body(doc,
        "短弹幕处理：对于\"111\"、\"...\"、\"？\"等无意义短弹幕，"
        "系统通过正则表达式进行识别，"
        "并推测其可能的意图（如\"可能在刷屏\"或\"一脸问号\"），"
        "将推测意图附加到发送给LLM的提示中，"
        "帮助模型生成更有针对性的回复。")
    add_body(doc,
        "规则降级：当DeepSeek API不可用（网络故障、API密钥未配置等）时，"
        "系统自动回退到基于关键词匹配的规则引擎，"
        "使用预设的回复模板生成简单的互动回复，"
        "保证直播间的基本互动不会中断。")
    add_body(doc,
        "待回复队列：当AI正在生成回复时，新到达的弹幕会被缓存到待回复队列中，"
        "等待当前回复完成后处理最新的一条，"
        "确保AI始终回复最新的弹幕内容。")
    add_body(doc,
        "此外，系统还在 AI 层加入了\"神秘观众\"过滤规则——平台分配的"
        "\"神秘观众***\"用户ID不参与计数、不写入历史、也不触发回复——避免"
        "占用对话窗口的位置而干扰真实观众的互动。该规则属于业务过滤，"
        "在 capture 层不做任何修改，只在 AI 层判定。")
    add_figure(doc, os.path.join(FIG_DIR, "fig4_3_ai_decision.png"),
               "图4-3 AI 回复决策流程图")

    add_h2(doc, "4.4 语音合成模块设计")
    add_body(doc,
        "语音合成模块围绕VoiceManager（语音管理器）和TTSWorker（TTS工作器）"
        "两个核心组件构建。")
    add_body(doc,
        "VoiceManager负责管理语音供应商的配置和凭据，"
        "支持阿里云百炼等供应商。"
        "每个语音角色（主播、助播）关联一个声音条目（VoiceEntry），"
        "声音条目记录了样本路径、克隆状态和云端声音ID等信息。"
        "VoiceManager还负责合成缓存的管理，"
        "通过参数哈希避免重复调用API。")
    add_body(doc,
        "TTSWorker是一个基于asyncio.PriorityQueue的异步TTS队列处理器。"
        "当AI回复生成后，回复文本被封装为TTS任务并入队。"
        "TTSWorker按优先级出队任务，调用VoiceManager的合成接口生成音频，"
        "然后通过本地音频播放器进行播放。"
        "队列支持中断（高优先级任务可以中断正在播放的低优先级音频）和超时保护。")
    add_body(doc,
        "音频播放方面，系统在Windows平台使用winsound的异步播放模式，"
        "在其他平台使用pygame.mixer，确保音频播放不会阻塞主线程。"
        "在程序退出阶段，系统会显式调用 stop_all_audio() 主动中断仍在异步播放的"
        "winsound 任务，避免出现\"窗口已关闭但仍在出声\"的尴尬状态。")

    add_h2(doc, "4.5 模块依赖与关键参数")
    add_body(doc,
        "为帮助理解系统的模块组织方式，图4-4给出了主要模块之间的依赖与调用关系。"
        "其中 ui.py 与 CaptureWorker 形成 UI/逻辑的分界线，"
        "CaptureWorker 作为唯一的协调中心负责对外部服务的访问。")
    add_figure(doc, os.path.join(FIG_DIR, "fig4_4_module_deps.png"),
               "图4-4 模块依赖关系图")
    add_body(doc, "表4-2列出了系统中影响互动节奏与稳定性的关键参数。")
    add_table_caption(doc, "表4-2 系统关键参数表")
    add_table(doc,
        headers=["参数名称", "取值", "设计作用"],
        rows=[
            ["reply_interval", "30 秒", "限制两次自动回复的最小间隔，避免 AI 刷屏。"],
            ["USER_COOLDOWN_SEC", "60 秒", "限制同一用户重复触发自动回复。"],
            ["_CHAT_SIGNATURE_WINDOW_SEC", "2 秒",
             "对 WS 与 DOM 重复弹幕按签名去重。"],
            ["_CHAT_FUSION_WINDOW_SEC", "3 秒",
             "为 WS 和 DOM 通道预留互补融合窗口。"],
            ["obs_hit_window", "10 秒",
             "OBS 关键词命中计数滑动窗口。"],
            ["obs_rule_cooldown", "60 秒",
             "单条 OBS 规则触发后进入冷却的时长。"],
            ["llm_timeout", "6 至 12 秒",
             "DeepSeek 调用的自适应超时上限，超时即触发规则降级。"],
            ["hls_segment_time", "2 秒",
             "ffmpeg 切片时长，配合 3 个 playlist 长度以平衡延迟与稳定性。"],
        ],
        col_widths_cm=[5.0, 3.0, 7.5])

    # CHAPTER 5
    add_h1(doc, "5 系统实现")
    add_h2(doc, "5.1 登录会话管理与弹幕抓取")
    add_body(doc,
        "登录会话管理模块利用Playwright的launch_persistent_context函数"
        "创建持久化浏览器上下文，将Cookie和localStorage等浏览器数据"
        "存储在本地browser_data目录中。用户首次使用时，"
        "系统打开抖音首页并等待扫码登录（最长2.5分钟），"
        "登录成功后浏览器状态自动保存，后续启动时无需重新扫码。")
    add_body(doc,
        "系统通过检查sessionid Cookie和页面URL来判断登录状态。"
        "如果检测到登录过期（页面被重定向到passport.douyin.com），"
        "系统会提示用户重新扫码。")
    add_body(doc,
        "弹幕抓取模块在RoomCapture类中实现。启动时，"
        "系统在已认证的浏览器上下文中创建新页面，"
        "注入WebSocket引用捕获脚本和DOM观察脚本，"
        "然后导航到直播间URL并自动刷新以确保WebSocket连接稳定建立。"
        "页面还配置了媒体资源拦截规则（block media类型请求），"
        "减少视频流带来的带宽消耗，只保留弹幕数据传输所需的WebSocket连接。"
        "登录与抓取的实现流程如图5-1所示。")
    add_figure(doc, os.path.join(FIG_DIR, "fig5_1_login_flow.png"),
               "图5-1 登录与弹幕抓取实现流程图")
    add_body(doc,
        "图5-2是连接直播间后的直播页运行截图，"
        "页面左侧展示按消息类型过滤后的实时弹幕，右侧为 AI 回复区。"
        "页面顶部的状态栏可分别启用/停用弹幕抓取和 AI 自动回复，"
        "对应 CaptureWorker 中的两个独立开关。")
    add_figure(doc, os.path.join(FIG_DIR, "ui_live.png"),
               "图5-2 直播页界面截图", width_inch=5.6)

    add_h2(doc, "5.2 协议解码与消息处理")
    add_body(doc,
        "协议解码在DanmakuDecoder类中实现，采用无状态设计（decode方法无副作用）。"
        "原始二进制帧经过三层解码：首先解析PushFrame信封获取gzip压缩的payload，"
        "然后解压缩得到Response消息体，最后遍历Response中的消息列表，"
        "根据method字段分派到对应的类型解析器"
        "（ChatMessage、GiftMessage、LikeMessage等）。")
    add_body(doc,
        "系统维护了方法统计（method_counts）和解析失败计数（parse_fail_counts），"
        "每60秒输出一次统计日志，帮助开发者监控协议变化和异常情况。"
        "对于未知的WebSocket方法，系统会记录首次观测的样本数据"
        "到debug_payload.log文件中，便于后续分析。"
        "解码与统计的整体路径如图5-3所示。")
    add_body(doc,
        "消息处理采用WS/DOM双通道融合策略。对于chat类型消息，"
        "WebSocket通道为优先通道，DOM通道为补充通道。"
        "系统通过（昵称, 内容）标准化的chat signature进行去重，"
        "在2秒窗口内同一签名只保留最早到达的消息。"
        "去重统计指标（ws_chat_total、ws_chat_complete、dom_chat_promoted等）"
        "每60秒记录一次。")
    add_figure(doc, os.path.join(FIG_DIR, "fig5_3_protocol_decode.png"),
               "图5-3 协议解码与消息处理流程图")

    add_h2(doc, "5.3 AI 智能回复引擎实现")
    add_body(doc,
        "AI智能回复引擎在AIReplyEngine类中实现。当新弹幕到达时，"
        "引擎首先检查自动回复开关、屏蔽词、全局回复间隔和用户冷却等前置条件。"
        "通过检查后，引擎获取asyncio.Lock（确保同一时间只有一个回复在生成），"
        "构建包含系统提示词和对话历史的消息列表，调用DeepSeek API生成回复。")
    add_body(doc,
        "对话历史管理采用滑动窗口策略，保留最近4条对话记录"
        "（2条用户消息+2条助手回复），既提供了基本的上下文理解能力，"
        "又控制了API调用的token消耗。系统还集成了用户记忆管理（MemoryManager），"
        "将每次互动保存为JSONL文件，为未来的长期记忆功能预留了接口。")
    add_body(doc,
        "OBS联动模块在ObsActionController类中实现，"
        "通过关键词匹配和命中计数触发场景切换。"
        "每条弹幕内容经过casefold标准化后与规则库中的关键词进行匹配，"
        "匹配到的规则进入命中桶（hit bucket），"
        "在设定的时间窗口（默认10秒）内累计命中次数达到阈值（默认2次）时触发动作。"
        "动作执行后进入规则冷却期（默认60秒），防止同一规则被频繁触发。"
        "播放期间可以选择忽略新的触发请求。"
        "图5-4是 AI 配置页的运行截图，"
        "图中展示了人设维度配置区、节流冷却参数区以及调试入口。")
    add_figure(doc, os.path.join(FIG_DIR, "ui_ai_config.png"),
               "图5-4 AI 配置页界面截图", width_inch=5.6)

    add_h2(doc, "5.4 OBS 联动与数字人推流")
    add_body(doc,
        "OBS联动模块以 ObsActionController 为核心。"
        "每条规则记录关键词、最小命中次数、计数窗口与冷却时长，"
        "图5-5给出了一次典型触发的时序示意：弹幕命中两次关键词后"
        "切换到指定场景，规则进入 60 秒冷却，"
        "并在媒体播放完成后自动切回主场景。")
    add_figure(doc, os.path.join(FIG_DIR, "fig5_7_obs_timing.png"),
               "图5-5 OBS 关键词触发动作时序示意图", width_inch=5.6)
    add_body(doc,
        "数字人推流模块在DigitalHumanPipeline类中实现，采用多步骤流水线架构："
        "首先检查ffmpeg可用性，然后通过VoiceManager合成主播话术的语音，"
        "接着启动ffmpeg进程将绿幕视频和TTS音频合成为HLS切片流，"
        "同时启动本地HTTP服务器提供HLS文件访问，"
        "最后通过OBS WebSocket自动配置Media Source和色度键滤镜。"
        "整体流水线如图5-6所示。")
    add_body(doc,
        "流水线的每个步骤都有独立的状态标识"
        "（IDLE、SYNTHESIZING、STARTING_SERVER、CONFIGURING_OBS、"
        "PUSHING、STREAMING、STOPPING、ERROR、CANCELLED），"
        "便于UI层实时显示进度。停止操作会依次终止ffmpeg进程、关闭HTTP服务器、"
        "清理临时文件并移除OBS中的媒体源。")
    add_body(doc,
        "ffmpeg视频合成采用HLS协议，设置2秒的切片时长和3个切片的播放列表大小，"
        "启用delete_segments标志自动清理过期切片，避免磁盘空间持续增长。"
        "视频编码使用libx264（ultrafast preset），音频编码使用AAC，"
        "兼顾了编码速度和画面质量。")
    add_figure(doc, os.path.join(FIG_DIR, "fig5_5_dh_pipeline.png"),
               "图5-6 数字人推流流水线状态图", width_inch=5.6)
    add_body(doc,
        "图5-7是 OBS 联动页与数字人推流页的运行截图。"
        "上半区为 OBS 连接状态与规则编辑器，"
        "下半区为数字人推流的视频/语音/状态配置入口。")
    add_figure(doc, os.path.join(FIG_DIR, "ui_obs.png"),
               "图5-7 OBS 联动页界面截图", width_inch=5.6)
    add_figure(doc, os.path.join(FIG_DIR, "ui_digital_human.png"),
               "图5-8 数字人推流页界面截图", width_inch=5.6)

    add_h2(doc, "5.5 开发问题与解决方案")
    add_body(doc,
        "在为期数月的实际开发过程中，本系统遭遇了多个具有代表性的工程问题。"
        "下面按照\"场景 - 现象 - 排查 - 方案\"的顺序，"
        "分别记录这些问题的诊断与处理过程，"
        "并在表5-1中给出整体汇总。")
    add_body(doc,
        "（1）扫码登录反复失效。在早期版本中，每次重启 Aiszr 都需要重新扫码登录抖音，"
        "严重影响日常使用体验。"
        "经过排查发现，Playwright 默认每次启动都会创建一个全新的临时浏览器目录，"
        "Cookie 与 localStorage 都无法跨进程保留。"
        "解决方案是改用 Playwright 的 launch_persistent_context 接口，"
        "把浏览器数据落地到项目根目录下的 browser_data/ 文件夹；"
        "并在启动时通过检查 sessionid Cookie 与页面 URL"
        "判断登录是否已经失效，失效时再触发一次扫码流程。"
        "改造后扫码次数从\"每次启动一次\"降为\"每隔约 30 天一次\"。")
    add_body(doc,
        "（2）WebSocket 弹幕昵称缺失。在某次抖音前端发版后，"
        "部分 chat 消息的 user 字段开始出现仅有 user_id、缺失 nick_name 的情况。"
        "由于 AI 回复需要昵称构造称呼，这些缺字段消息会被丢弃，"
        "造成可观察的回复漏发。"
        "排查时同时打开 WS 通道日志和 DOM 注入脚本日志后发现，"
        "页面 DOM 仍然能渲染出完整的昵称——只是 WS 协议本身不再下发。"
        "解决方案是在原有 WS 通道之外，"
        "用 MutationObserver 监听聊天区域的新增节点并通过 expose_binding 把"
        "（昵称，内容）回传给 Python 端；"
        "随后在 capture 层加入 2 秒签名窗口对两通道结果做去重和补全融合。"
        "改造后 chat 消息的昵称完整率从约 88% 提升至接近 100%。")
    add_body(doc,
        "（3）未知协议方法难以定位。抖音不定期会新增 WebSocket method（"
        "例如新版的礼物连击通知、虚拟陪伴提示等），"
        "新版本上线后 DanmakuDecoder 会在大量\"未识别 method\"上抛出异常，"
        "但日志里只有一条简短的错误信息，无法用于复盘。"
        "解决方案是在 decoder 中增加 method_counts、parse_fail_counts 两个计数器，"
        "每 60 秒输出一次方法分布和失败次数；"
        "对于首次出现的未知 method，把样本写入 debug_payload.log，"
        "包含十六进制原帧与解压后的负载结构。"
        "这种\"先记录后补丁\"的方式可以让协议适配工作从\"线上灭火\""
        "退化为\"离线追加\"。")
    add_body(doc,
        "（4）AI 回复过于频繁。最早接入 DeepSeek 时采用\"每条弹幕都送 LLM\"的"
        "天真策略，结果在弹幕高峰期 AI 回复会以接近秒级的频率刷屏，"
        "反而把直播间真实观众的互动挤出可视区域。"
        "解决方案分为三层：第一层加入 30 秒全局最小回复间隔，"
        "确保两次自动回复至少间隔半分钟；"
        "第二层针对每个用户 ID 单独维护 60 秒冷却窗口，"
        "防止同一用户连续提问占据全部回复机会；"
        "第三层把 LLM 调用封装在 asyncio.Lock 中串行执行，"
        "避免并行请求把 token 配额烧光。"
        "三层叠加后实际单位时间内的回复条数下降了一个数量级，"
        "但每一条都更精准。")
    add_body(doc,
        "（5）ffmpeg HLS 切片在 OBS 端有明显卡顿。在数字人推流的早期实现里，"
        "把 hls_time 设置为 1 秒、playlist 长度设置为 2 之后，"
        "OBS Media Source 会出现明显的\"跳一跳\"现象，"
        "节奏感非常糟糕。"
        "排查发现 ffmpeg 在每次切片生成的瞬间会有一段缓冲空窗期，"
        "切片越短、空窗期占比越大，OBS 缓存又来不及补齐。"
        "经过多组对照试验后，把 hls_time 调到 2 秒，"
        "把 hls_list_size 调到 3，"
        "并保留 delete_segments 标志自动清理过期切片。"
        "调整后画面持续平滑，磁盘占用维持在 30 MB 以内。")
    add_body(doc,
        "（6）winsound 异步播放在程序关闭时仍在出声。Windows 平台为了避免阻塞主线程"
        "采用了 winsound.SND_ASYNC 异步播放，"
        "但这一模式在窗口关闭、Qt 退出后并不会立即停止——"
        "用户在 UI 上看到的是\"程序已经退出\"，但喇叭还在念上一条 AI 回复，"
        "体验非常诡异。"
        "解决方案是在 AiszrApp 的 closeEvent 与 QApplication.aboutToQuit"
        "两条路径上都显式调用 audio_output.stop_all_audio()，"
        "底层封装 winsound.PlaySound(None, 0) 立即中断所有 SND_ASYNC 任务，"
        "并同步关停 pygame.mixer 的备用通道。"
        "改造后\"窗口已关但语音继续\"的问题彻底消失。")
    add_body(doc,
        "（7）PyQt5 信号跨线程更新 UI 偶发卡死。CaptureWorker 跑在独立 QThread 中，"
        "asyncio 协程产生的弹幕事件需要回到 Qt 主线程才能写入界面。"
        "早期通过手工调用 QMetaObject.invokeMethod 完成跨线程调度，"
        "在压测时偶尔会出现死锁，原因是 Worker 线程持有的 asyncio.Lock 与"
        "主线程的 Qt 事件锁形成了交叉等待。"
        "解决方案是把所有跨线程触发改为 PyQt Signal/Slot，"
        "并通过 asyncio.run_coroutine_threadsafe 把主线程发起的协程提交回 Worker 的"
        "事件循环。这条规则在团队内被总结为：\"asyncio 的事归 asyncio，"
        "Qt 的事归 Qt，跨线程只通过 Signal 与 run_coroutine_threadsafe 沟通。\"")

    add_table_caption(doc, "表5-1 开发问题与解决方案汇总")
    add_table(doc,
        headers=["问题", "产生原因", "解决方案"],
        rows=[
            ["扫码登录反复失效", "默认临时上下文不保存 Cookie",
             "改用 launch_persistent_context，结合 sessionid 检查复用登录态。"],
            ["WS 弹幕昵称缺失", "前端协议字段变更",
             "加入 DOM MutationObserver 补偿，并通过 2 秒签名窗口去重融合。"],
            ["未知协议方法难定位", "平台新增消息类型",
             "记录 method_counts/parse_fail_counts，并把样本写入 debug_payload.log。"],
            ["AI 回复过于频繁", "对每条弹幕同步调用 LLM",
             "30 秒全局节流 + 60 秒用户冷却 + asyncio.Lock 串行生成。"],
            ["HLS 切片卡顿", "切片过短引起缓冲空窗",
             "hls_time 调至 2 秒、playlist 长度 3，配合 delete_segments 清理。"],
            ["winsound 关闭后仍出声", "SND_ASYNC 不随窗口退出",
             "closeEvent / aboutToQuit 中显式 stop_all_audio() 中断全部音频。"],
            ["跨线程 UI 偶发卡死", "asyncio.Lock 与 Qt 事件锁交叉等待",
             "全部跨线程交互改为 Signal + run_coroutine_threadsafe。"],
        ],
        col_widths_cm=[3.5, 4.2, 7.8])

    # CHAPTER 6
    add_h1(doc, "6 系统测试")
    add_h2(doc, "6.1 功能测试")
    add_body(doc, "功能测试覆盖了系统的各个核心模块，测试结果如表6-1所示。")
    add_table_caption(doc, "表6-1 功能测试结果")
    add_table(doc,
        headers=["测试模块", "测试内容", "测试结果"],
        rows=[
            ["登录会话", "扫码登录、会话持久化、过期检测", "通过"],
            ["弹幕抓取", "WS 拦截、DOM 观察、双通道去重", "通过"],
            ["协议解码", "PushFrame/Response/ChatMessage 解析", "通过"],
            ["AI 回复", "DeepSeek 接口调用、人设回复、节流冷却", "通过"],
            ["OBS 联动", "关键词命中、规则冷却、播完切回主场景", "通过"],
            ["语音合成", "声音克隆、TTS 合成、缓存命中", "通过"],
            ["数字人推流", "ffmpeg 启停、HLS 切片、OBS 媒体源配置", "通过"],
        ],
        col_widths_cm=[4.0, 7.5, 4.0])

    add_h2(doc, "6.2 性能测试")
    add_body(doc,
        "性能测试主要关注弹幕抓取的实时性和AI回复的响应速度。"
        "测试环境为Windows 11系统，Intel i5处理器，16GB内存，"
        "使用pytest框架执行测试。")
    add_body(doc,
        "在弹幕解码性能方面，DanmakuDecoder的decode方法处理单帧数据的"
        "平均耗时在1毫秒以内，能够满足高频率弹幕场景下的实时解码需求。"
        "协议解码的稳定性经过多次长时间运行测试验证，"
        "在4小时连续运行中未出现内存泄漏或性能退化。"
        "图6-1左侧为 4 小时压测前 60 分钟的 CPU 占用曲线，"
        "右侧为同期内存（RSS）变化，两者均维持在预期区间。")
    add_figure(doc, os.path.join(FIG_DIR, "fig5_8_metrics.png"),
               "图6-1 系统资源占用观测（前 60 分钟）", width_inch=5.6)
    add_body(doc,
        "在AI回复性能方面，DeepSeek API的平均响应时间为3至5秒（100 token输出），"
        "加上本地处理开销，从弹幕接收到回复展示的端到端延迟控制在6至8秒以内，"
        "满足直播互动的时效性要求。"
        "图6-2展示了从弹幕到达到 TTS 播放完成的逐阶段耗时拆分，"
        "以及一次 4 小时压测中弹幕来源构成。")
    add_figure(doc, os.path.join(FIG_DIR, "fig6_1_perf.png"),
               "图6-2 AI 回复链路阶段耗时与弹幕来源分布", width_inch=5.6)
    add_body(doc,
        "系统整体资源占用合理，CPU使用率在空闲时低于5%，"
        "在弹幕高峰时低于15%，内存占用稳定在200MB左右，"
        "适合在普通配置的Windows电脑上运行。")

    # 结语
    add_h1(doc, "结　语")
    add_body(doc,
        "本文设计并实现了一个基于大语言模型的直播互动智能助手系统（Aiszr），"
        "围绕\"稳定抓弹幕 -> AI 文本回复 -> OBS 动作联动 -> AI 语音播报\""
        "的完整闭环，为直播电商场景提供了一套综合性智能化解决方案。")
    add_body(doc, "系统的主要创新点和贡献包括：")
    add_body(doc,
        "（1）提出了基于Playwright的被动式WebSocket拦截与DOM观察的双通道弹幕采集方案，"
        "无需逆向协议签名即可稳定抓取直播弹幕，在协议变更时具有更好的适应性。")
    add_body(doc,
        "（2）设计了包含人设系统、回复节流、用户冷却、短弹幕意图推测和规则降级的"
        "多层AI回复策略，实现了高质量、低干扰的直播间智能互动。")
    add_body(doc,
        "（3）集成了弹幕抓取、AI回复、OBS联动、语音播报和数字人推流等多种能力，"
        "形成完整的直播辅助闭环，主播无需在多个工具之间频繁切换。")
    add_body(doc,
        "系统已通过功能测试和性能测试验证，各模块运行稳定可靠。"
        "但本系统仍存在一些不足之处：语音克隆依赖阿里云百炼的单供应商服务，"
        "未来需要扩展多供应商支持；弹幕抓取目前仅支持抖音平台，"
        "需要适配更多直播平台；"
        "AI回复的上下文理解能力受限于滑动窗口策略，"
        "需要引入更高效的长期记忆机制。")
    add_body(doc,
        "未来的研究方向包括：扩展对快手、淘宝直播等平台的支持；"
        "引入RAG（检索增强生成）技术提升AI回复的商品知识准确性；"
        "研究巨量百应平台API实现文字场控发送能力；"
        "优化多房间并发管理策略，支持同时监控多个直播间。")

    # 参考文献
    add_h1(doc, "参考文献")
    refs = [
        "[1] 深度求索. DeepSeek API 文档[EB/OL]. https://platform.deepseek.com/api-docs, 2025.",
        "[2] 微软. Playwright Python 文档[EB/OL]. https://playwright.dev/python/, 2025.",
        "[3] 阿里云. DashScope 语音合成 API 文档[EB/OL]. https://help.aliyun.com/document_detail/2712536.html, 2025.",
        "[4] OBS Project. OBS WebSocket v5 Protocol[EB/OL]. https://github.com/obsproject/obs-websocket/blob/master/docs/generated/protocol.md, 2024.",
        "[5] Google. Protocol Buffers Documentation[EB/OL]. https://protobuf.dev/, 2025.",
        "[6] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems, 2017: 5998-6008.",
        "[7] zhonghangAlex. DySpider: 抖音直播弹幕抓取[EB/OL]. https://github.com/zhonghangAlex/DySpider, 2024.",
        "[8] OpenAI. GPT-4 Technical Report[J]. arXiv preprint arXiv:2303.08774, 2023.",
        "[9] Touvron H, Lavril T, Izacard G, et al. LLaMA: Open and Efficient Foundation Language Models[J]. arXiv preprint arXiv:2302.13971, 2023.",
        "[10] FastAPI. FastAPI Framework Documentation[EB/OL]. https://fastapi.tiangolo.com/, 2025.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(r)
        set_run_cn_font(run, FONT_SONG, 12)

    # 致谢
    add_h1(doc, "致　谢")
    add_body(doc,
        "经过数月的学习和努力，本次毕业设计已经接近尾声。"
        "在此，我要衷心感谢所有给予我帮助和支持的人。")
    add_body(doc,
        "首先，我要特别感谢我的指导老师章文老师。"
        "在整个毕业设计过程中，章老师从选题方向、技术路线到论文撰写，"
        "都给予了我耐心的指导和悉心的帮助。"
        "每当我在技术实现上遇到困难时，章老师总能给出建设性的建议，"
        "帮助我理清思路、找到解决方案。"
        "章老师严谨的治学态度和精益求精的工作作风，让我受益匪浅。")
    add_body(doc,
        "其次，我要感谢设计与信息学院的各位任课老师。"
        "在三年的人工智能技术应用专业学习中，"
        "是他们教授了我Python编程、数据结构、机器学习、深度学习等专业基础知识，"
        "为本次毕业设计的顺利完成奠定了坚实的基础。")
    add_body(doc,
        "同时，我还要感谢我的同学们。"
        "在开发过程中，我们互相交流技术方案、讨论遇到的问题，"
        "他们的建议和帮助对我启发很大。"
        "特别感谢在系统测试阶段给予我反馈和建议的朋友们。")
    add_body(doc,
        "最后，我要感谢我的家人，是他们在背后默默的支持和鼓励，"
        "让我能够专注于学业和毕业设计，顺利完成大学阶段的学习任务。")
    add_body(doc,
        "在未来的工作和学习中，我将继续努力，"
        "将所学知识应用于实际工作中，不断提升自己的技术能力和专业素养。")

    doc.save(OUT_PATH)
    print(f"saved -> {OUT_PATH}")
    print(f"size  -> {os.path.getsize(OUT_PATH):,} bytes")


if __name__ == "__main__":
    build()
